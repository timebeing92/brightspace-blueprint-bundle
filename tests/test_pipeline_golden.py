"""Golden test: a fresh run over examples/sample_export.zip must reproduce the
committed worked example in examples/sample_course__blueprint_bundle/."""
from __future__ import annotations

import json
from pathlib import Path

import pytest

from conftest import BUNDLE_ROOT, EXPECTED_BUNDLE

TEXT_ARTIFACTS = [
    "README.md",
    "sample_course__blueprint.md",
    "sample_course__cps_blueprint.md",
    "sample_course__course_structure.md",
    "sample_course__course_activities.md",
    "sample_course__course_qa.md",
    "sample_course__docx_structure.md",
    "sample_export__inventory.md",
    "sample_export__manifest_probe.md",
]

JSON_ARTIFACTS = [
    "sample_course__blueprint.json",
    "sample_course__course_structure.json",
    "sample_course__course_activities.json",
    "sample_course__course_qa.json",
    "sample_course__docx_structure.json",
    "sample_course__rubrics.json",
    "sample_export__inventory.json",
    "sample_export__manifest_probe.json",
]

BINARY_ARTIFACTS = [
    "sample_course__blueprint.docx",
    "sample_course__course_activities.xlsx",
    "sample_course__rubrics.xlsx",
]


def test_pipeline_exits_cleanly(golden_run):
    assert golden_run.proc.returncode == 0, (
        f"pipeline failed\nSTDOUT:\n{golden_run.proc.stdout}\n"
        f"STDERR:\n{golden_run.proc.stderr}"
    )


def test_bundle_contains_expected_files(golden_run):
    expected = {p.name for p in EXPECTED_BUNDLE.iterdir() if p.is_file()}
    produced = {p.name for p in golden_run.bundle_dir.iterdir() if p.is_file()}
    assert produced == expected


@pytest.mark.parametrize("name", TEXT_ARTIFACTS)
def test_text_artifact_matches_golden(golden_run, name):
    produced = (golden_run.bundle_dir / name).read_text(encoding="utf-8")
    expected = (EXPECTED_BUNDLE / name).read_text(encoding="utf-8")
    assert produced == expected


@pytest.mark.parametrize("name", JSON_ARTIFACTS)
def test_json_artifact_matches_golden(golden_run, name):
    produced = json.loads((golden_run.bundle_dir / name).read_text(encoding="utf-8"))
    expected = json.loads((EXPECTED_BUNDLE / name).read_text(encoding="utf-8"))
    assert produced == expected


@pytest.mark.parametrize("name", BINARY_ARTIFACTS)
def test_binary_artifact_present_and_nontrivial(golden_run, name):
    path = golden_run.bundle_dir / name
    assert path.is_file()
    assert path.stat().st_size > 1024


def _docx_visible_text(path: Path) -> list[str]:
    from docx import Document

    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(p.text for p in cell.paragraphs)
    return lines


def test_docx_text_matches_golden(golden_run):
    produced = _docx_visible_text(golden_run.bundle_dir / "sample_course__blueprint.docx")
    expected = _docx_visible_text(EXPECTED_BUNDLE / "sample_course__blueprint.docx")
    assert produced == expected


def _workbook_cells(path: Path) -> dict[str, list[list[object]]]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True)
    return {
        ws.title: [[cell.value for cell in row] for row in ws.iter_rows()]
        for ws in wb.worksheets
    }


def test_workbook_matches_golden(golden_run):
    produced = _workbook_cells(golden_run.bundle_dir / "sample_course__course_activities.xlsx")
    expected = _workbook_cells(EXPECTED_BUNDLE / "sample_course__course_activities.xlsx")
    assert produced == expected


def test_blueprint_json_validates_against_schema(golden_run):
    jsonschema = pytest.importorskip("jsonschema")
    schema = json.loads(
        (BUNDLE_ROOT / "schemas" / "blueprint_schema.json").read_text(encoding="utf-8")
    )
    model = json.loads(
        (golden_run.bundle_dir / "sample_course__blueprint.json").read_text(encoding="utf-8")
    )
    jsonschema.validate(model, schema)


def test_rubrics_json_validates_against_schema(golden_run):
    jsonschema = pytest.importorskip("jsonschema")
    schema = json.loads(
        (BUNDLE_ROOT / "schemas" / "rubrics_schema.json").read_text(encoding="utf-8")
    )
    model = json.loads(
        (golden_run.bundle_dir / "sample_course__rubrics.json").read_text(encoding="utf-8")
    )
    jsonschema.validate(model, schema)
    assert len(model["rubrics"]) == 1
    assert len(model["rubrics"][0]["levels"]) == 2
    assert len(model["rubrics"][0]["criteria"]) == 2
