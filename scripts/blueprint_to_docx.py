#!/usr/bin/env python3
"""Render a structured blueprint model (JSON) to DOCX.

Input is the ``<label>__blueprint.json`` produced by build_blueprint_bundle.py.
The model contract is documented in ``schemas/blueprint_schema.json``. The output
uses a stable course-blueprint review structure:

- course header line + course title
- single-column front-matter tables (Description / Materials / Course Learning Outcomes)
- "Course Introduction" heading + intro text + "Course Content:"
- one full-width table per week, using either section labels above content
  rows (default) or section labels in a left column:
    Overview
    Learning Objectives
    Assigned Reading and Multimedia
    Assignment(s) and Instructions
    Discussion Board Prompts
    Checklist                        (only when present)
    Other course sections            (only when present)

In the default layout, scaffold labels sit in shaded rows immediately above
their extracted content. In the optional left-column layout, the same labels sit
in shaded left cells and the content remains in full-page-width tables.
Bullets use the document's native List Bullet styles (real hanging indents,
nesting by level) when the style base defines them.

When a template DOCX is supplied with ``--template``, it is opened as a generic
style base so the output can inherit local fonts, heading styles, and page setup.
Its body content is cleared and regenerated from the model.

Requires python-docx (``pip install python-docx``). If it is missing this exits
with code 3 and a clear message; the Markdown blueprint is unaffected.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover - exercised only without the dependency
    sys.stderr.write(
        "python-docx is not installed; cannot render DOCX.\n"
        "Install it with:  pip install python-docx\n"
        "(The Markdown blueprint was still produced.)\n"
    )
    raise SystemExit(3)

LIVE_SCHEMES = ("http://", "https://", "mailto:")

NOT_FOUND_FIELD = "Needs review: not found in export extraction."
NOT_FOUND_LIST = "None found in export extraction."
LABEL_FALLBACK = "(blank in source - fill in during review)"

# Stable row labels for the blueprint review document.
OVERVIEW_LABEL = "Overview: (add an introduction to the week's topic and activities here, with references as needed)"
OBJECTIVES_LABEL = (
    "Learning Objectives: Must follow the guidelines in this Learning Objectives Guide.\n"
    "Students will be able to:"
)
ASSIGNMENTS_LABEL = "Assignment(s) and Instructions:"
DISCUSSIONS_LABEL = "Discussion Board Prompts:"
READING_LABEL = "Assigned Reading and Multimedia: (add links, articles, textbook readings, videos). Include style-correct citations."
CHECKLIST_LABEL = "Checklist (mirrored from the export)"
OTHER_LABEL = "Other course sections (mirrored from the export)"
BEFORE_WEEK_LABEL = "Before Week 1: Additional Resources and Information"

DESCRIPTION_LABEL = (
    "COURSE DESCRIPTION (keep in mind the Course Description must match the published catalog, "
    "any changes must be approved by the Program and planned in advance)"
)
MATERIALS_LABEL = "TEXTBOOK/S OR REQUIRED MATERIALS"
OUTCOMES_LABEL = "COURSE LEARNING OUTCOMES"


# --------------------------------------------------------------------------- #
# Style helpers
# --------------------------------------------------------------------------- #
def clean_label(value: str) -> str:
    """Normalize a display label before the renderer adds its own trailing colon."""
    return clean_text(value).rstrip(":").strip()


def clean_text(value: str) -> str:
    return " ".join(xml_safe_text(value).split())


def xml_safe_text(value: str) -> str:
    """Remove characters that cannot appear in WordprocessingML XML text."""
    return "".join(
        char if is_xml_compatible_char(char) else " "
        for char in str(value or "")
    )


def is_xml_compatible_char(char: str) -> bool:
    code = ord(char)
    return (
        code in (0x09, 0x0A, 0x0D)
        or 0x20 <= code <= 0xD7FF
        or 0xE000 <= code <= 0xFFFD
        or 0x10000 <= code <= 0x10FFFF
    )


def style_or_none(doc: "Document", name: str):
    """Return a style by name if the document defines it, else None."""
    try:
        return doc.styles[name]
    except KeyError:
        return None


def apply_heading(doc: "Document", paragraph, level_name: str) -> None:
    style = style_or_none(doc, level_name)
    if style is not None:
        paragraph.style = style


def _apply_para_style(paragraph, style_name: str) -> bool:
    """Apply a named paragraph style if the document defines it; report success."""
    try:
        paragraph.style = style_name
        return True
    except KeyError:
        return False


def _space(paragraph, *, before: int | None = None, after: int | None = None) -> None:
    """Set paragraph spacing in points (template Normal packs paragraphs tight)."""
    fmt = paragraph.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)


def _twips(inches: float) -> int:
    return int(round(inches * 1440))


def usable_width_inches(doc: "Document") -> float:
    """Return the writable page width for the document's first section."""
    section = doc.sections[0]
    return section.page_width.inches - section.left_margin.inches - section.right_margin.inches


def _replace_child(parent, tag: str, child) -> None:
    for existing in parent.findall(qn(tag)):
        parent.remove(existing)
    parent.append(child)


def set_column_widths(table, widths_inches: tuple[float, ...]) -> None:
    """Fix table, grid, and cell widths so Word expands it to the writable page width."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(_twips(sum(widths_inches))))
    tbl_w.set(qn("w:type"), "dxa")
    _replace_child(tbl_pr, "w:tblW", tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    _replace_child(tbl_pr, "w:tblInd", tbl_ind)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(_twips(width)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(_twips(width)))
            tc_w.set(qn("w:type"), "dxa")
            _replace_child(tc_pr, "w:tcW", tc_w)


def set_cell_margins(table, *, top: int = 60, bottom: int = 60,
                     start: int = 110, end: int = 110) -> None:
    """Give every cell in the table a little padding (units: dxa, 20ths of a point)."""
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tbl_pr.append(margins)


def _trim_leading_empty(cell) -> None:
    """Drop the cell's default empty first paragraph once content follows it."""
    paragraphs = cell.paragraphs
    if (
        len(paragraphs) > 1
        and not paragraphs[0].runs
        and not paragraphs[0]._p.findall(qn("w:hyperlink"))
    ):
        paragraphs[0]._p.getparent().remove(paragraphs[0]._p)


def set_cell_background(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_paragraph_background(paragraph, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def add_table_borders(table) -> None:
    """Add single black borders to every edge/inside line of a table."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tbl_pr.append(borders)


# --------------------------------------------------------------------------- #
# Cell content helpers
# --------------------------------------------------------------------------- #
def first_paragraph(cell):
    """Return the cell's initial empty paragraph (reuse it, don't append a blank)."""
    return cell.paragraphs[0]


def write_label(cell, text: str) -> None:
    """Write a template scaffold label into its own shaded cell, in smaller
    bold type so the extracted content in the neighboring cell stands out."""
    para = first_paragraph(cell)
    for chunk in text.split("\n"):
        if para.runs:
            para = cell.add_paragraph()
        _space(para, before=0, after=2)
        run = para.add_run(chunk)
        run.bold = True
        run.font.size = Pt(9)
    set_cell_background(cell, "F2F2F2")


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Append a real, clickable hyperlink run to a paragraph (cell or body)."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = xml_safe_text(text)
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _emit_runs(paragraph, runs: list[dict]) -> None:
    for run in runs:
        text = xml_safe_text(run.get("text", ""))
        if not text:
            continue
        href = xml_safe_text(run.get("href") or "").strip()
        if href.startswith(LIVE_SCHEMES):
            add_hyperlink(paragraph, href, text)
        elif href:
            paragraph.add_run(text)
            note = paragraph.add_run(f" ({href})")
            note.font.size = Pt(8)
            note.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            paragraph.add_run(text)


def _write_missing(container, missing: str) -> None:
    para = container.add_paragraph()
    run = para.add_run(missing)
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


BULLET_STYLE_BY_LEVEL = {1: "List Bullet", 2: "List Bullet 2", 3: "List Bullet 3"}

_BULLET_NUM_ID_CACHE: dict[int, str | None] = {}


def _find_bullet_num_id(doc_part) -> str | None:
    """Find a numbering definition whose first level is a bullet; None when the
    document has no numbering part."""
    try:
        numbering = doc_part.part_related_by(RELATIONSHIP_TYPE.NUMBERING).element
    except (KeyError, AttributeError):
        return None
    bullet_abstract_ids = set()
    for abstract in numbering.findall(qn("w:abstractNum")):
        lvl0 = next(
            (lvl for lvl in abstract.findall(qn("w:lvl")) if lvl.get(qn("w:ilvl")) == "0"),
            None,
        )
        fmt = lvl0.find(qn("w:numFmt")) if lvl0 is not None else None
        if fmt is not None and fmt.get(qn("w:val")) == "bullet":
            bullet_abstract_ids.add(abstract.get(qn("w:abstractNumId")))
    for num in numbering.findall(qn("w:num")):
        ref = num.find(qn("w:abstractNumId"))
        if ref is not None and ref.get(qn("w:val")) in bullet_abstract_ids:
            return num.get(qn("w:numId"))
    return None


def _bullet_num_id(paragraph) -> str | None:
    """Memoized bullet numId lookup for the paragraph's document."""
    key = id(paragraph.part)
    if key not in _BULLET_NUM_ID_CACHE:
        _BULLET_NUM_ID_CACHE[key] = _find_bullet_num_id(paragraph.part)
    return _BULLET_NUM_ID_CACHE[key]


def _apply_native_bullet(paragraph, num_id: str, level: int) -> None:
    """Attach real Word list numbering (numPr) to a fresh paragraph."""
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 9) - 1))
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), num_id)
    numpr.append(ilvl)
    numpr.append(numid)
    ppr.insert(0, numpr)


def _emit_divider(container, *, before: int = 6, after: int = 6) -> None:
    para = container.add_paragraph()
    _space(para, before=before, after=after)
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9D9D9")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _emit_block(container, block: dict, *, previous_kind: str = "", wrapper_fill: str = "") -> None:
    """Add one block as a paragraph; only list items get a bullet, paragraphs
    stay prose. Bullets become native Word lists: List Bullet styles when the
    style base defines them, else a numPr reference to the template's own
    bullet numbering, else a manual hanging-indent bullet."""
    kind = block.get("kind")
    if kind == "divider":
        _emit_divider(container)
        return

    runs = [r for r in block.get("runs", []) if r.get("text")]
    if not runs:
        return
    para = container.add_paragraph()
    if kind == "label":
        before = 2 if previous_kind in {"section_label", "label"} else 8
        _space(para, before=before, after=2)
        if wrapper_fill:
            set_paragraph_background(para, wrapper_fill)
        for run in runs:
            text = clean_text(run.get("text", ""))
            if text:
                para.add_run(text).bold = True
        return
    if kind == "practice":
        _space(para, before=6, after=5)
        set_paragraph_background(para, wrapper_fill or "EEF7F2")
        title = clean_text(runs[0].get("text", ""))
        href = clean_text(runs[0].get("href", ""))
        title_run = para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(9)
        if href:
            note = para.add_run(f" ({href})")
            note.font.size = Pt(8)
            note.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        summary = clean_text(block.get("meta", {}).get("summary", ""))
        if summary:
            para.add_run().add_break()
            summary_run = para.add_run(summary)
            summary_run.font.size = Pt(8)
            summary_run.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)
        return
    if kind in {"visual", "dropdown", "embed", "image", "file", "hidden"}:
        label = {
            "visual": "Visual cue: ",
            "dropdown": "Dropdown / expandable section: ",
            "embed": f"{clean_text(block.get('meta', {}).get('embed_type', '')) or 'Embedded media'}: ",
            "image": "Embedded image: ",
            "file": "Attached course file: ",
            "hidden": "Hidden manifest item: ",
        }[kind]
        _space(para, before=6, after=4)
        fill = {
            "visual": "F2F6FA",
            "dropdown": "F4F2FA",
            "embed": "F7F9FB",
            "image": "F7F9FB",
            "file": "F7F9FB",
            "hidden": "FFF4E5",
        }[kind]
        set_paragraph_background(para, wrapper_fill or fill)
        prefix = para.add_run(label)
        prefix.bold = True
        prefix.font.size = Pt(9)
        prefix.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)
        _emit_runs(para, runs)
        if kind == "visual" and block.get("blocks"):
            previous_child_kind = "visual"
            for child in block["blocks"]:
                _emit_block(container, child, previous_kind=previous_child_kind, wrapper_fill=fill)
                previous_child_kind = child.get("kind", "")
        return
    if kind == "li":
        level = max(1, int(block.get("level") or 1))
        if not _apply_para_style(para, BULLET_STYLE_BY_LEVEL.get(min(level, 3), "List Bullet")):
            num_id = _bullet_num_id(para)
            if num_id is not None:
                _apply_native_bullet(para, num_id, level)
            else:
                fmt = para.paragraph_format
                fmt.left_indent = Pt(18 * level)
                fmt.first_line_indent = Pt(-9)
                para.add_run("• ")
        _space(para, before=0, after=2)
    else:
        _space(para, before=0, after=6)
    if wrapper_fill:
        set_paragraph_background(para, wrapper_fill)
    _emit_runs(para, runs)


def write_blocks(container, blocks: list[dict], *, missing: str) -> None:
    """Render blocks (paragraphs / bullets with link runs) into a cell or the body."""
    if not blocks:
        _write_missing(container, missing)
        return
    previous_kind = ""
    for block in blocks:
        _emit_block(container, block, previous_kind=previous_kind)
        previous_kind = block.get("kind", "")


def _add_section_divider(container) -> None:
    _emit_divider(container)


def should_divide_labeled_sections(previous: dict | None, current: dict, divider: bool | str) -> bool:
    if not previous or not divider:
        return False
    if divider == "page":
        previous_page = clean_text(previous.get("source_page", ""))
        current_page = clean_text(current.get("source_page", ""))
        if previous_page and current_page:
            return previous_page != current_page
        return False
    if divider == "object":
        return True
    return True


def write_value_labeled(container, sections: list[dict], *, missing: str, divider: bool | str = False) -> None:
    """Write labeled sections: a bold label line, then its blocks (kind-aware).
    Sections after the first get breathing room so entries don't run together."""
    if not sections:
        _write_missing(container, missing)
        return
    previous_section: dict | None = None
    for index, section in enumerate(sections):
        has_divider = should_divide_labeled_sections(previous_section, section, divider)
        if has_divider:
            _add_section_divider(container)
        label = clean_label(section.get("label", ""))
        if label:
            para = container.add_paragraph()
            _space(para, before=2 if has_divider else (8 if index else 0), after=2)
            para.add_run(f"{label}:").bold = True
        previous_kind = "section_label" if label else ""
        for block in section.get("blocks", []):
            _emit_block(container, block, previous_kind=previous_kind)
            previous_kind = block.get("kind", "")
        previous_section = section


# --------------------------------------------------------------------------- #
# Body construction
# --------------------------------------------------------------------------- #
def clear_body(doc: "Document") -> None:
    """Remove all paragraphs and tables, preserving the final sectPr."""
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def add_front_matter_table(doc: "Document", label: str, blocks: list[dict]) -> None:
    table = doc.add_table(rows=2, cols=1)
    add_table_borders(table)
    set_column_widths(table, (usable_width_inches(doc),))
    set_cell_margins(table)
    write_label(table.cell(0, 0), label)
    value = table.cell(1, 0)
    write_blocks(value, blocks, missing=NOT_FOUND_FIELD)
    _trim_leading_empty(value)
    doc.add_paragraph()


def week_section_rows(week: dict):
    # (label, fill) section pairs. Due day-of-week rides along in the
    # assignment text; coded numeric dates are term-relative and not encoded.
    rows = [
        (OVERVIEW_LABEL,
         lambda cell: write_blocks(cell, week.get("overview", []), missing=NOT_FOUND_FIELD)),
        (OBJECTIVES_LABEL,
         lambda cell: write_blocks(cell, week.get("learning_objectives", []), missing=NOT_FOUND_FIELD)),
        (READING_LABEL,
         lambda cell: write_value_labeled(cell, week.get("resources", []), missing=NOT_FOUND_LIST)),
        (ASSIGNMENTS_LABEL,
         lambda cell: write_value_labeled(
             cell, week.get("assignments", []), missing=NOT_FOUND_LIST, divider="object"
         )),
        (DISCUSSIONS_LABEL,
         lambda cell: write_value_labeled(
             cell, week.get("discussions", []), missing=NOT_FOUND_LIST, divider="object"
         )),
    ]
    if week.get("checklist"):
        rows.append((CHECKLIST_LABEL,
                     lambda cell: write_value_labeled(cell, week["checklist"], missing=NOT_FOUND_LIST)))
    if week.get("other_sections"):
        rows.append((OTHER_LABEL,
                     lambda cell: write_value_labeled(
                         cell,
                         week["other_sections"],
                         missing=NOT_FOUND_LIST,
                         divider="page",
                     )))
    return rows


def add_week_table_top(doc: "Document", week: dict) -> None:
    rows = week_section_rows(week)

    table = doc.add_table(rows=len(rows) * 2, cols=1)
    add_table_borders(table)
    set_column_widths(table, (usable_width_inches(doc),))
    set_cell_margins(table)
    for index, (label, fill) in enumerate(rows):
        write_label(table.cell(index * 2, 0), label)
        value = table.cell(index * 2 + 1, 0)
        fill(value)
        _trim_leading_empty(value)

    doc.add_paragraph()


def add_week_table_left(doc: "Document", week: dict) -> None:
    rows = week_section_rows(week)
    full_width = usable_width_inches(doc)
    label_width = min(1.8, max(1.35, full_width * 0.28))
    value_width = full_width - label_width

    table = doc.add_table(rows=len(rows), cols=2)
    add_table_borders(table)
    set_column_widths(table, (label_width, value_width))
    set_cell_margins(table)
    for index, (label, fill) in enumerate(rows):
        write_label(table.cell(index, 0), label)
        value = table.cell(index, 1)
        fill(value)
        _trim_leading_empty(value)

    doc.add_paragraph()


def add_week_table(doc: "Document", week: dict, *, section_layout: str = "top") -> None:
    apply_heading(doc, doc.add_paragraph(week.get("title", "Course Module")), "Heading 2")
    if section_layout == "left":
        add_week_table_left(doc, week)
    else:
        add_week_table_top(doc, week)


def add_before_week_table(doc: "Document", sections: list[dict]) -> None:
    apply_heading(doc, doc.add_paragraph(BEFORE_WEEK_LABEL), "Heading 2")
    table = doc.add_table(rows=len(sections) * 2, cols=1)
    add_table_borders(table)
    set_column_widths(table, (usable_width_inches(doc),))
    set_cell_margins(table)
    for index, section in enumerate(sections):
        write_label(table.cell(index * 2, 0), clean_label(section.get("label", "")) or BEFORE_WEEK_LABEL)
        value = table.cell(index * 2 + 1, 0)
        write_blocks(value, section.get("blocks", []), missing=NOT_FOUND_LIST)
        _trim_leading_empty(value)
    doc.add_paragraph()


def add_simple_section(doc: "Document", heading: str, items: list[str]) -> None:
    apply_heading(doc, doc.add_paragraph(heading), "Heading 2")
    if not items:
        doc.add_paragraph("None.")
        return
    for item in items:
        _emit_block(doc, {"kind": "li", "level": 1, "runs": [{"text": item, "href": ""}]})


def render(model: dict, doc: "Document", *, section_layout: str = "top") -> None:
    clear_body(doc)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(model.get("course_title") or "Course Blueprint")
    title_run.bold = True
    title_run.font.size = Pt(16)

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Evidence mode: extracted text is source-derived; missing fields remain marked for review."
    )
    note_run.italic = True
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

    fm = model.get("front_matter", {})
    add_front_matter_table(doc, DESCRIPTION_LABEL, fm.get("course_description", []))
    add_front_matter_table(doc, MATERIALS_LABEL, fm.get("required_materials", []))
    add_front_matter_table(doc, OUTCOMES_LABEL, fm.get("course_learning_outcomes", []))

    apply_heading(doc, doc.add_paragraph("Course Introduction"), "Heading 2")
    write_blocks(doc, fm.get("course_introduction", []), missing=NOT_FOUND_FIELD)

    before_week = model.get("before_week_1", [])
    if before_week:
        add_before_week_table(doc, before_week)

    content_para = doc.add_paragraph()
    content_para.add_run("Course Content:").bold = True
    doc.add_paragraph()

    for week in model.get("weeks", []):
        add_week_table(doc, week, section_layout=section_layout)

    unplaced = model.get("unplaced_activities", {})
    if unplaced.get("assignments") or unplaced.get("discussions"):
        apply_heading(doc, doc.add_paragraph("Unplaced Activities"), "Heading 2")
        doc.add_paragraph(
            "These activities were extracted from D2L object XML but were not connected to a "
            "module quicklink by resource code."
        )
        if unplaced.get("assignments"):
            apply_heading(doc, doc.add_paragraph("Assignments"), "Heading 4")
            write_value_labeled(doc, unplaced["assignments"], missing=NOT_FOUND_LIST, divider="object")
        if unplaced.get("discussions"):
            apply_heading(doc, doc.add_paragraph("Discussions"), "Heading 4")
            write_value_labeled(doc, unplaced["discussions"], missing=NOT_FOUND_LIST, divider="object")

    add_simple_section(doc, "Extraction Notes", model.get("diagnostics", []) or ["None."])


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("model", type=Path, help="Path to <label>__blueprint.json")
    parser.add_argument("--output", type=Path, required=True, help="Output .docx path")
    parser.add_argument(
        "--template",
        type=Path,
        default=None,
        help="Optional .docx to use as a style base; body content is regenerated",
    )
    parser.add_argument(
        "--section-layout",
        choices=("top", "left"),
        default="top",
        help="DOCX weekly section label layout: shaded top rows (default) or shaded left column.",
    )
    args = parser.parse_args(argv)

    if not args.model.exists():
        raise SystemExit(f"error: model not found: {args.model}")
    model = json.loads(args.model.read_text(encoding="utf-8"))

    if args.template and args.template.exists():
        doc = Document(str(args.template))
    else:
        if args.template:
            sys.stderr.write(f"note: template not found ({args.template}); building DOCX from scratch.\n")
        doc = Document()

    render(model, doc, section_layout=args.section_layout)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"docx: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
