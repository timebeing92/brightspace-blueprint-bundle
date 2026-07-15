#!/usr/bin/env python3
"""Render coursecraft.rubrics/1 JSON to a reviewer-friendly DOCX.

The same renderer is used for the standalone ``<label>__rubrics.docx`` artifact
and for the Rubric Appendix inside the main blueprint DOCX. The input is the
canonical JSON emitted by ``extract_rubrics_to_workbook.py --json``.
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover - exercised only without the dependency
    sys.stderr.write(
        "python-docx is not installed; cannot render rubric DOCX.\n"
        "Install it with:  pip install python-docx\n"
    )
    raise SystemExit(3)


def clean(value: object) -> str:
    return " ".join(str(value or "").split())


def style_or_none(doc: "Document", name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return None


def apply_heading(doc: "Document", paragraph, level_name: str) -> None:
    style = style_or_none(doc, level_name)
    if style is not None:
        paragraph.style = style


def _space(paragraph, *, before: int | None = None, after: int | None = None) -> None:
    fmt = paragraph.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)


def _twips(inches: float) -> int:
    return int(round(inches * 1440))


def usable_width_inches(doc: "Document") -> float:
    section = doc.sections[0]
    return section.page_width.inches - section.left_margin.inches - section.right_margin.inches


def _replace_child(parent, tag: str, child) -> None:
    for existing in parent.findall(qn(tag)):
        parent.remove(existing)
    parent.append(child)


def set_column_widths(table, widths_inches: tuple[float, ...]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(_twips(sum(widths_inches))))
    tbl_w.set(qn("w:type"), "dxa")
    _replace_child(tbl_pr, "w:tblW", tbl_w)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(_twips(width)))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = OxmlElement("w:tcW")
            tc_w.set(qn("w:w"), str(_twips(width)))
            tc_w.set(qn("w:type"), "dxa")
            _replace_child(tc_pr, "w:tcW", tc_w)


def set_cell_margins(table, *, top: int = 70, bottom: int = 70,
                     start: int = 100, end: int = 100) -> None:
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tbl_pr.append(margins)


def set_cell_background(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tbl_pr.append(borders)


def format_points(value: object) -> str:
    text = clean(value)
    if not text:
        return ""
    try:
        num = float(text)
    except ValueError:
        return text
    if num.is_integer():
        return str(int(num))
    return f"{num:g}"


def header_cell(cell, text: str) -> None:
    para = cell.paragraphs[0]
    _space(para, before=0, after=0)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    set_cell_background(cell, "F2F2F2")


def write_small_note(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def rubric_usage_map(activities_model: dict | None) -> dict[str, list[str]]:
    """Return rubric id -> activity labels from course_activities JSON."""
    usage: dict[str, list[str]] = {}
    if not activities_model:
        return usage

    def add(rubric_id: str, label: str) -> None:
        rubric_id = clean(rubric_id)
        label = clean(label)
        if not rubric_id or not label:
            return
        usage.setdefault(rubric_id, [])
        if label not in usage[rubric_id]:
            usage[rubric_id].append(label)

    for folder in activities_model.get("dropbox_folders", []) or []:
        name = clean(folder.get("name")) or clean(folder.get("resource_code")) or "Assignment"
        grade = clean(folder.get("grade_item_name"))
        suffix = f" (grade item: {grade})" if grade and grade != name else ""
        for rubric_id in folder.get("rubric_ids", []) or []:
            add(rubric_id, f"Assignment: {name}{suffix}")

    for topic in activities_model.get("discussions", []) or []:
        if clean(topic.get("kind")) == "forum":
            continue
        name = clean(topic.get("title")) or clean(topic.get("resource_code")) or "Discussion"
        grade = clean(topic.get("grade_item_name"))
        suffix = f" (grade item: {grade})" if grade and grade != name else ""
        for rubric_id in topic.get("rubric_ids", []) or []:
            add(rubric_id, f"Discussion: {name}{suffix}")

    for quiz in activities_model.get("quizzes", []) or []:
        name = clean(quiz.get("title")) or clean(quiz.get("resource_code")) or "Quiz"
        grade = clean(quiz.get("grade_item_name"))
        suffix = f" (grade item: {grade})" if grade and grade != name else ""
        for rubric_id in quiz.get("rubric_ids", []) or []:
            add(rubric_id, f"Quiz: {name}{suffix}")

    return usage


def rubric_table_count(rubrics_model: dict | None) -> int:
    if not rubrics_model:
        return 0
    count = 0
    for rubric in rubrics_model.get("rubrics", []) or []:
        count += 1
        if rubric.get("overall_levels"):
            count += 1
    return count


def add_rubric_grid_table(doc: "Document", rubric: dict) -> None:
    levels = rubric.get("levels", []) or []
    criteria = rubric.get("criteria", []) or []
    cols = max(2, len(levels) + 1)
    rows = max(1, len(criteria)) + 1
    table = doc.add_table(rows=rows, cols=cols)
    add_table_borders(table)
    set_cell_margins(table)

    full_width = usable_width_inches(doc)
    criteria_width = min(2.0, max(1.35, full_width * 0.28))
    level_width = (full_width - criteria_width) / max(1, cols - 1)
    set_column_widths(table, (criteria_width, *([level_width] * (cols - 1))))

    header_cell(table.cell(0, 0), "Criterion")
    for index, level in enumerate(levels, start=1):
        name = clean(level.get("name")) or f"Level {index}"
        band = clean(level.get("score_band"))
        header_cell(table.cell(0, index), f"{name}\n{band}" if band else name)

    if not criteria:
        cell = table.cell(1, 0)
        para = cell.paragraphs[0]
        para.add_run("No criteria were extracted.").italic = True
        return

    for row_index, criterion in enumerate(criteria, start=1):
        name_cell = table.cell(row_index, 0)
        para = name_cell.paragraphs[0]
        _space(para, before=0, after=0)
        run = para.add_run(clean(criterion.get("name")) or "Criterion")
        run.bold = True
        run.font.size = Pt(9)

        cells_by_level = {
            clean(cell.get("level_id")): cell
            for cell in criterion.get("cells", []) or []
        }
        for col_index, level in enumerate(levels, start=1):
            level_id = clean(level.get("level_id"))
            rubric_cell = cells_by_level.get(level_id, {})
            cell = table.cell(row_index, col_index)
            para = cell.paragraphs[0]
            _space(para, before=0, after=0)
            points = format_points(rubric_cell.get("points"))
            if points:
                points_run = para.add_run(f"{points} pts")
                points_run.bold = True
                points_run.font.size = Pt(8)
            description = clean(rubric_cell.get("description"))
            if description:
                if points:
                    para.add_run().add_break()
                desc_run = para.add_run(description)
                desc_run.font.size = Pt(8)

    doc.add_paragraph()


def add_overall_levels_table(doc: "Document", rubric: dict) -> None:
    overall_levels = rubric.get("overall_levels", []) or []
    if not overall_levels:
        return
    para = doc.add_paragraph()
    _space(para, before=4, after=2)
    para.add_run("Overall scoring bands").bold = True

    table = doc.add_table(rows=len(overall_levels) + 1, cols=3)
    add_table_borders(table)
    set_cell_margins(table)
    full_width = usable_width_inches(doc)
    set_column_widths(table, (full_width * 0.24, full_width * 0.18, full_width * 0.58))
    for col, text in enumerate(("Band", "Starts at", "Description / feedback")):
        header_cell(table.cell(0, col), text)

    for row_index, level in enumerate(overall_levels, start=1):
        values = [
            clean(level.get("name")) or "Overall level",
            format_points(level.get("range_start_value")),
            clean(level.get("description")),
        ]
        feedback = clean(level.get("feedback"))
        if feedback:
            values[2] = f"{values[2]} Feedback: {feedback}".strip()
        for col, text in enumerate(values):
            para = table.cell(row_index, col).paragraphs[0]
            _space(para, before=0, after=0)
            run = para.add_run(text)
            run.font.size = Pt(8)
    doc.add_paragraph()


def render_rubrics_section(
    doc: "Document",
    rubrics_model: dict | None,
    activities_model: dict | None = None,
    *,
    title: str = "Rubric Appendix",
    page_break: bool = False,
) -> int:
    rubrics = (rubrics_model or {}).get("rubrics", []) or []
    if not rubrics:
        return 0

    if page_break:
        doc.add_page_break()
    apply_heading(doc, doc.add_paragraph(title), "Heading 2")
    note = doc.add_paragraph()
    write_small_note(
        note,
        "Rubric grids are source-derived from rubrics_d2l.xml. Activity usage is "
        "listed when the export exposes a rubric association.",
    )

    usage = rubric_usage_map(activities_model)
    tables_added = 0
    for index, rubric in enumerate(rubrics):
        if index:
            doc.add_paragraph()
        name = clean(rubric.get("name")) or "Untitled Rubric"
        apply_heading(doc, doc.add_paragraph(name), "Heading 3")

        metadata_bits = []
        if clean(rubric.get("id")):
            metadata_bits.append(f"Rubric ID: {clean(rubric.get('id'))}")
        if clean(rubric.get("resource_code")):
            metadata_bits.append(f"Resource code: {clean(rubric.get('resource_code'))}")
        if clean(rubric.get("scoring_method")):
            metadata_bits.append(f"Scoring method: {clean(rubric.get('scoring_method'))}")
        if metadata_bits:
            meta = doc.add_paragraph()
            write_small_note(meta, " | ".join(metadata_bits))

        used_by = usage.get(clean(rubric.get("id")), [])
        if used_by:
            used_para = doc.add_paragraph()
            used_para.add_run("Used by: ").bold = True
            used_para.add_run("; ".join(used_by))

        description = clean(rubric.get("description"))
        if description:
            desc_para = doc.add_paragraph()
            desc_para.add_run(description)

        add_rubric_grid_table(doc, rubric)
        tables_added += 1
        if rubric.get("overall_levels"):
            add_overall_levels_table(doc, rubric)
            tables_added += 1

    return tables_added


def render_standalone(
    rubrics_model: dict,
    doc: "Document",
    activities_model: dict | None = None,
    *,
    title: str = "Rubrics",
) -> None:
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    render_rubrics_section(
        doc,
        rubrics_model,
        activities_model,
        title="Rubric Grids",
        page_break=False,
    )


def load_json(path: Path | None) -> dict | None:
    if path is None:
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("rubrics_json", type=Path, help="Path to <label>__rubrics.json")
    parser.add_argument("--output", type=Path, required=True, help="Output .docx path")
    parser.add_argument(
        "--activities-json",
        type=Path,
        default=None,
        help="Optional <label>__course_activities.json for Used by lines",
    )
    parser.add_argument("--title", default="Rubrics", help="Document title")
    args = parser.parse_args(argv)

    if not args.rubrics_json.exists():
        raise SystemExit(f"error: rubric JSON not found: {args.rubrics_json}")
    rubrics_model = load_json(args.rubrics_json)
    activities_model = load_json(args.activities_json) if args.activities_json and args.activities_json.exists() else None

    doc = Document()
    render_standalone(rubrics_model or {}, doc, activities_model, title=args.title)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output))
    print(f"rubrics docx: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
