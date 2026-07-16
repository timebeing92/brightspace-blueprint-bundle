#!/usr/bin/env python3
"""Structural QA for a rendered blueprint DOCX — no LibreOffice required.

Reads the generated ``<label>__blueprint.docx`` back and verifies, against the
``<label>__blueprint.json`` model it was rendered from, that the document's
structure survived: the package parses, every relationship reference resolves,
every hyperlink points at a URL the model actually contains, the expected
tables exist with the expected shape, and the course/week titles are present.

This is the lightweight default check; the LibreOffice/Poppler visual render
QA (``render_blueprint_docx.py``) remains the optional pixel-level deep pass.

Vocabulary matches the course QA report: breaks (structural defects), warnings
(suspicious but viewable), notes (stats). Exit codes: 0 = no breaks,
1 = breaks found, 2 = usage error.
"""
from __future__ import annotations

import argparse
import json
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover - exercised only without the dependency
    sys.stderr.write(
        "python-docx is not installed; cannot inspect the DOCX.\n"
        "Install it with:  pip install python-docx\n"
    )
    raise SystemExit(3)

# Must mirror blueprint_to_docx.py: only these schemes become live hyperlinks.
LIVE_SCHEMES = ("http://", "https://", "mailto:")
# Renderer sections per week table: overview, objectives, reading,
# assignments, discussions (+ checklist / other_sections when present).
BASE_WEEK_SECTIONS = 5
FRONT_MATTER_TABLES = 3

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def xml_safe_text(value: str) -> str:
    """Inline copy of the bundle's common_xml.xml_safe_text (this checker is
    mirrored into repos whose common_xml does not carry the helper)."""
    def compatible(char: str) -> bool:
        code = ord(char)
        return (
            code in (0x09, 0x0A, 0x0D)
            or 0x20 <= code <= 0xD7FF
            or 0xE000 <= code <= 0xFFFD
            or 0x10000 <= code <= 0x10FFFF
        )
    return "".join(char if compatible(char) else " " for char in str(value or ""))


# --------------------------------------------------------------------------- #
# Model side: what the renderer should have emitted
# --------------------------------------------------------------------------- #
def _walk_blocks(blocks):
    """Yield live-hyperlink hrefs from blocks, mirroring the renderer's
    _emit_block/_emit_runs traversal: label and practice blocks render their
    hrefs as plain text; only visual blocks recurse into children."""
    for block in blocks or []:
        kind = block.get("kind")
        if kind in {"divider", "label", "practice"}:
            continue
        for run in block.get("runs", []):
            text = xml_safe_text(run.get("text", ""))
            if not text:
                continue
            href = xml_safe_text(run.get("href") or "").strip()
            if href.startswith(LIVE_SCHEMES):
                yield href
        if kind == "visual":
            yield from _walk_blocks(block.get("blocks"))


def _walk_sections(sections):
    for section in sections or []:
        yield from _walk_blocks(section.get("blocks"))


def model_link_hrefs(model: dict) -> list[str]:
    hrefs: list[str] = []
    front_matter = model.get("front_matter", {})
    for key in (
        "course_description",
        "required_materials",
        "course_learning_outcomes",
        "course_introduction",
    ):
        hrefs.extend(_walk_blocks(front_matter.get(key)))
    hrefs.extend(_walk_sections(model.get("before_week_1")))
    for week in model.get("weeks", []):
        hrefs.extend(_walk_blocks(week.get("overview")))
        hrefs.extend(_walk_blocks(week.get("learning_objectives")))
        for key in ("resources", "assignments", "discussions", "checklist", "other_sections"):
            hrefs.extend(_walk_sections(week.get(key)))
    unplaced = model.get("unplaced_activities", {})
    hrefs.extend(_walk_sections(unplaced.get("assignments")))
    hrefs.extend(_walk_sections(unplaced.get("discussions")))
    return hrefs


def expected_week_shape(week: dict, layout: str) -> tuple[int, int]:
    sections = BASE_WEEK_SECTIONS
    if week.get("checklist"):
        sections += 1
    if week.get("other_sections"):
        sections += 1
    if layout == "left":
        return sections, 2
    return sections * 2, 1


# --------------------------------------------------------------------------- #
# DOCX side
# --------------------------------------------------------------------------- #
def read_relationships(docx_path: Path) -> tuple[dict[str, str], list[tuple[str, str]]]:
    """Return ({rel id: target}, [(element localname, referenced rel id), ...])."""
    with zipfile.ZipFile(docx_path) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        rels_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
    rels = {rel.get("Id"): rel.get("Target", "") for rel in rels_root}
    referenced: list[tuple[str, str]] = []
    for element in document.iter():
        for attr in (f"{{{NS_R}}}id", f"{{{NS_R}}}embed"):
            rel_id = element.get(attr)
            if rel_id:
                referenced.append((element.tag.split("}", 1)[-1], rel_id))
    return rels, referenced


def visible_text_lines(doc) -> list[str]:
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                lines.extend(p.text for p in cell.paragraphs)
    return lines


def normalized_visible_text(value: object) -> str:
    """Mirror the renderer's whitespace cleanup for heading comparisons."""
    return " ".join(xml_safe_text(str(value or "")).split())


# --------------------------------------------------------------------------- #
# The check
# --------------------------------------------------------------------------- #
def rubric_table_count(rubrics_model: dict | None) -> int:
    if not rubrics_model:
        return 0
    count = 0
    for rubric in rubrics_model.get("rubrics", []) or []:
        count += 1
        if rubric.get("overall_levels"):
            count += 1
    return count


def check(docx_path: Path, model: dict, layout: str, rubrics_model: dict | None = None) -> dict:
    breaks: list[str] = []
    warnings: list[str] = []
    notes: list[str] = []

    try:
        doc = Document(str(docx_path))
        rels, referenced = read_relationships(docx_path)
    except Exception as exc:  # damaged package: nothing else is checkable
        return {
            "breaks": [f"DOCX did not open as a valid package: {exc}"],
            "warnings": [],
            "notes": [],
            "stats": {},
        }

    # Every r:id / r:embed in the document must resolve to a relationship.
    dangling = [(tag, rid) for tag, rid in referenced if rid not in rels]
    for tag, rid in dangling[:10]:
        breaks.append(f"Dangling relationship reference: <{tag}> points at missing id {rid}")
    if len(dangling) > 10:
        breaks.append(f"...and {len(dangling) - 10} more dangling relationship references")

    # Every hyperlink target must be a URL the model contains.
    model_hrefs = model_link_hrefs(model)
    model_href_set = {href.strip() for href in model_hrefs}
    docx_targets = [
        rels[rid]
        for tag, rid in referenced
        if tag == "hyperlink" and rid in rels
    ]
    foreign = [t for t in docx_targets if t.strip() not in model_href_set]
    for target in foreign[:10]:
        breaks.append(f"Hyperlink target not present in the model: {target}")
    if len(foreign) > 10:
        breaks.append(f"...and {len(foreign) - 10} more foreign hyperlink targets")
    if len(docx_targets) != len(model_hrefs):
        warnings.append(
            f"Hyperlink count mismatch: model has {len(model_hrefs)} live links, "
            f"DOCX has {len(docx_targets)}"
        )

    # Table census: front matter + optional before-week + one per week.
    weeks = model.get("weeks", [])
    before_week_tables = 1 if model.get("before_week_1") else 0
    rubric_tables = rubric_table_count(rubrics_model)
    expected_tables = FRONT_MATTER_TABLES + before_week_tables + len(weeks) + rubric_tables
    actual_tables = len(doc.tables)
    if actual_tables != expected_tables:
        breaks.append(
            f"Table count mismatch: expected {expected_tables} "
            f"({FRONT_MATTER_TABLES} front matter + {before_week_tables} before-week "
            f"+ {len(weeks)} weeks + {rubric_tables} rubric), found {actual_tables}"
        )
    else:
        week_tables = doc.tables[FRONT_MATTER_TABLES + before_week_tables:]
        for week, table in zip(weeks, week_tables[:len(weeks)]):
            want_rows, want_cols = expected_week_shape(week, layout)
            got_rows, got_cols = len(table.rows), len(table.columns)
            if (got_rows, got_cols) != (want_rows, want_cols):
                warnings.append(
                    f"Week table shape for {week.get('title', '(untitled)')!r}: "
                    f"expected {want_rows}x{want_cols} ({layout} layout), got {got_rows}x{got_cols}"
                )

    # Titles must survive as body paragraphs.
    body_texts = {p.text for p in doc.paragraphs}
    normalized_body_texts = {normalized_visible_text(text) for text in body_texts}
    course_title = model.get("course_title") or "Course Blueprint"
    if normalized_visible_text(course_title) not in normalized_body_texts:
        breaks.append(f"Course title not found in the document: {course_title!r}")
    for week in weeks:
        title = week.get("title", "Course Module")
        if normalized_visible_text(title) not in normalized_body_texts:
            breaks.append(f"Week heading not found in the document: {title!r}")
    rubrics = (rubrics_model or {}).get("rubrics", []) or []
    if rubrics:
        if normalized_visible_text("Rubric Appendix") not in normalized_body_texts:
            breaks.append("Rubric Appendix heading not found in the document")
        for rubric in rubrics:
            name = rubric.get("name") or "Untitled Rubric"
            if normalized_visible_text(name) not in normalized_body_texts:
                breaks.append(f"Rubric heading not found in the document: {name!r}")

    lines = visible_text_lines(doc)
    stats = {
        "hyperlinks": len(docx_targets),
        "model_live_links": len(model_hrefs),
        "tables": actual_tables,
        "rubric_tables": rubric_tables,
        "paragraphs": len(lines),
        "visible_characters": sum(len(line) for line in lines),
        "weeks": len(weeks),
        "layout": layout,
    }
    notes.append(f"{stats['hyperlinks']} hyperlinks, all resolving to relationships"
                 if not dangling and not foreign
                 else f"{stats['hyperlinks']} hyperlinks inspected")
    notes.append(f"{stats['tables']} tables, {stats['paragraphs']} paragraphs, "
                 f"{stats['visible_characters']} visible characters")
    notes.append(f"{stats['weeks']} week section(s), {layout!r} section layout")
    if rubric_tables:
        notes.append(f"{rubric_tables} rubric appendix table(s)")

    return {"breaks": breaks, "warnings": warnings, "notes": notes, "stats": stats}


# --------------------------------------------------------------------------- #
# Reports
# --------------------------------------------------------------------------- #
def markdown_report(label: str, docx_name: str, model_name: str, result: dict) -> str:
    lines = [
        f"# DOCX Structure QA — {label}",
        "",
        f"- Document: `{docx_name}`",
        f"- Model: `{model_name}`",
        f"- Breaks: **{len(result['breaks'])}** · Warnings: **{len(result['warnings'])}** "
        f"· Notes: **{len(result['notes'])}**",
        "",
        "## Breaks",
        "",
    ]
    lines.extend(f"- {item}" for item in result["breaks"]) if result["breaks"] else lines.append("- none")
    lines.extend(["", "## Warnings", ""])
    lines.extend(f"- {item}" for item in result["warnings"]) if result["warnings"] else lines.append("- none")
    lines.extend(["", "## Notes", ""])
    lines.extend(f"- {item}" for item in result["notes"]) if result["notes"] else lines.append("- none")
    return "\n".join(lines) + "\n"


def summary_report(label: str, result: dict, report_path: Path | None) -> str:
    stats = result.get("stats", {})
    lines = [
        f"DOCX structure — {label}: {len(result['breaks'])} breaks · "
        f"{len(result['warnings'])} warnings · {len(result['notes'])} notes.",
    ]
    if stats:
        lines.append(
            f"  {stats.get('hyperlinks', 0)} hyperlinks · {stats.get('tables', 0)} tables "
            f"· {stats.get('paragraphs', 0)} paragraphs"
        )
    for item in result["breaks"][:3]:
        lines.append(f"  BREAK: {item}")
    if report_path is not None:
        lines.append(f"  Full report: {report_path}")
    return "\n".join(lines)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("docx", type=Path, help="Path to <label>__blueprint.docx")
    parser.add_argument("--model", type=Path, required=True, help="Path to <label>__blueprint.json")
    parser.add_argument(
        "--section-layout",
        choices=("top", "left"),
        default="top",
        help="Layout the DOCX was rendered with (affects expected table shape)",
    )
    parser.add_argument("--output-dir", default=None, help="Optional directory for markdown/json outputs")
    parser.add_argument(
        "--print-full",
        action="store_true",
        help="Print the full markdown report to stdout (default: a short summary)",
    )
    parser.add_argument(
        "--rubrics-json",
        type=Path,
        default=None,
        help="Optional <label>__rubrics.json when the blueprint DOCX includes a rubric appendix",
    )
    args = parser.parse_args(argv)

    if not args.docx.exists():
        print(f"error: DOCX not found: {args.docx}", file=sys.stderr)
        return 2
    if not args.model.exists():
        print(f"error: model not found: {args.model}", file=sys.stderr)
        return 2
    model = json.loads(args.model.read_text(encoding="utf-8"))
    rubrics_model = None
    if args.rubrics_json:
        if not args.rubrics_json.exists():
            print(f"error: rubric JSON not found: {args.rubrics_json}", file=sys.stderr)
            return 2
        rubrics_model = json.loads(args.rubrics_json.read_text(encoding="utf-8"))

    result = check(args.docx, model, args.section_layout, rubrics_model)

    label = args.docx.stem
    if label.endswith("__blueprint"):
        label = label[: -len("__blueprint")]
    md = markdown_report(label, args.docx.name, args.model.name, result)

    md_path = None
    if args.output_dir:
        output_dir = Path(args.output_dir).expanduser().resolve()
        output_dir.mkdir(parents=True, exist_ok=True)
        payload = {
            "label": label,
            "docx": args.docx.name,
            "model": args.model.name,
            "breaks": result["breaks"],
            "warnings": result["warnings"],
            "notes": result["notes"],
            "stats": result["stats"],
        }
        (output_dir / f"{label}__docx_structure.json").write_text(
            json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )
        md_path = output_dir / f"{label}__docx_structure.md"
        md_path.write_text(md, encoding="utf-8")

    print(md if args.print_full else summary_report(label, result, md_path))
    return 1 if result["breaks"] else 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
